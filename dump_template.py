"""Full diagnostic of the Word template - dump everything to a text file and try rendering."""
from docx import Document
from docxtpl import DocxTemplate
import traceback

docx_path = r'C:\Users\Cris\Desktop\AP Workers\AP Mitarbeiter 2026\Vorlage_Lebenslauf Muster\Vorlage_Lebenslauf_Schweißer Python.docx'

doc = Document(docx_path)

with open('/tmp/template_dump.txt', 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        f.write(f"P{i}: {repr(p.text)}\n")
        if len(p.runs) > 1 and ('{%' in p.text or '{{' in p.text):
            runs_text = [r.text for r in p.runs]
            f.write(f"  RUNS: {runs_text}\n")
    
    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---\n")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    f.write(f"  T{ti}.R{ri}.C{ci}.P{pi}: {repr(p.text)}\n")
                    if len(p.runs) > 1 and ('{%' in p.text or '{{' in p.text):
                        runs_text = [r.text for r in p.runs]
                        f.write(f"    RUNS: {runs_text}\n")

print("Dump written to /tmp/template_dump.txt")

# Now try rendering
print("\n=== ATTEMPTING RENDER ===")
try:
    tpl = DocxTemplate(docx_path)
    tpl.render({
        "education": {"higher_education": [], "further_training": []},
        "employment_history": [],
        "job_type": "Sudor",
        "language_skills": [],
        "starttermin": "01.03.2026",
    })
    print("SUCCESS!")
except Exception as e:
    print(f"RENDER ERROR: {e}")
    traceback.print_exc()
