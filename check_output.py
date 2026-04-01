"""Full detailed dump of populated output."""
from docx import Document
import os

path = os.path.join(os.environ['TEMP'], 'Populated_CV Paluch Dariusz Jan.docx')
doc = Document(path)

for ti, table in enumerate(doc.tables):
    print(f"=== TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ===")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            texts = [p.text for p in cell.paragraphs if p.text.strip()]
            if texts:
                for t in texts:
                    print(f"  R{ri}.C{ci}: {t}")
    print()
