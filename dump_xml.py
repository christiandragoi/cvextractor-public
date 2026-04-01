"""Verify the populated output."""
from docx import Document

doc = Document(r'C:\tmp\test_populated.docx')

print('=== PARAGRAPHS ===')
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        print(f'  {t}')

print()
print('=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'  --- Table {ti} ---')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                t = p.text.strip()
                if t:
                    print(f'    R{ri}.C{ci}: {t}')
    print()
