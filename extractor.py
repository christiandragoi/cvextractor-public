import os
import json
import re
import fitz          # PyMuPDF
from docx import Document


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    return "\n".join(page.get_text() for page in doc)


def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    full_text = []

    # 1. Paragraphs (main body)
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text)

    # 2. Tables (crucial for CVs!)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)

    # 3. Headers and Footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    if p.text.strip(): full_text.append(p.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    if p.text.strip(): full_text.append(p.text)

    # 4. Text Boxes (floating shapes)
    from lxml import etree
    xml_content = doc._element.xml
    # We look for all <w:t> tags that might be hidden in text boxes or shapes
    # Namespace for wordprocessingML
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    tree = etree.fromstring(xml_content.encode("utf-8"))
    # Find all text elements. doc.paragraphs already covers most, 
    # but we can find those NOT in the main flow or in special objects.
    # However, a simpler way in python-docx:
    for shape in doc.part.element.xpath('.//w:txbxContent//w:p'):
        p_text = "".join(node.text for node in shape.xpath('.//w:t') if node.text)
        if p_text.strip():
            full_text.append(p_text)

    # 5. Images (OCR fallback if still empty)
    res = "\n".join(full_text)
    return res


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")


def _pdf_pages_as_images(pdf_path: str, dpi: int = 150) -> list[tuple[str, str]]:
    """
    Render each PDF page to a temporary PNG file.
    Returns list of (tmp_path, mime_type) tuples.
    Caller is responsible for deleting the temp files.
    """
    import tempfile
    doc = fitz.open(pdf_path)
    results = []
    mat = fitz.Matrix(dpi / 72, dpi / 72)   # scale factor
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        pix.save(tmp.name)
        tmp.close()
        results.append((tmp.name, "image/png"))
    return results


def extract_text_from_scanned_pdf_openai(
    pdf_path: str, api_key: str,
    base_url: str = "https://api.openai.com/v1",
    model: str = "gpt-4o"
) -> str:
    """OCR a scanned PDF by rendering pages and using OpenAI vision."""
    import base64
    from openai import OpenAI
    pages = _pdf_pages_as_images(pdf_path)
    client = OpenAI(api_key=api_key, base_url=base_url)
    parts = []
    try:
        for img_path, mime in pages:
            with open(img_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": (
                        "Extract ALL text from this document page exactly as it appears. "
                        "Include every field, date, name and number. Return plain text only."
                    )},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                ]}],
                max_tokens=2048,
            )
            parts.append(resp.choices[0].message.content or "")
    finally:
        for img_path, _ in pages:
            try:
                os.unlink(img_path)
            except OSError:
                pass
    return "\n\n".join(parts)


def extract_text_from_scanned_pdf_gemini(
    pdf_path: str, api_key: str, model: str = "gemini-2.0-flash"
) -> str:
    """OCR a scanned PDF by rendering pages and using Gemini vision."""
    from google import genai
    from google.genai import types
    pages = _pdf_pages_as_images(pdf_path)
    client = genai.Client(api_key=api_key)
    parts = []
    try:
        for img_path, mime in pages:
            with open(img_path, "rb") as f:
                img_bytes = f.read()
            resp = client.models.generate_content(
                model=model,
                contents=[
                    types.Part.from_bytes(data=img_bytes, mime_type=mime),
                    "Extract ALL text from this document page exactly as it appears. Return plain text only.",
                ],
            )
            parts.append(resp.text or "")
    finally:
        for img_path, _ in pages:
            try:
                os.unlink(img_path)
            except OSError:
                pass
    return "\n\n".join(parts)


# ── Prompt ────────────────────────────────────────────────────────────────────

SYSTEM_MSG = (
    "You are a professional CV parser. "
    "Always respond with ONLY valid JSON — no markdown, no explanation, no code fences. "
    "CRITICAL: If the document contains NO identifiable names or data (e.g. empty or nonsensical text), "
    "set the 'name' field to 'N/A' or an error message. "
    "NEVER hallucinate example names like 'Hans Müller' or 'Max Mustermann' or 'Mustermann, Max'."
)

def build_prompt(text: str) -> str:
    return f"""Extract ALL information from the CV below.
Return ONLY a raw JSON object (no ```json fence, no extra text).

Required JSON structure:
{{
  "name": "Last Name, First Name",
  "birth_date": "DD.MM.YYYY",
  "nationality": "string",
  "id_expiry": "DD.MM.YYYY or null",
  "residence_permit_expiry": "DD.MM.YYYY or null",
  "job_role": "Schweißer | Elektriker | Schlosser | Lackierer | Mechaniker | etc.",
  "employment_history": [
    {{
      "employer":   "Company name",
      "position":   "Job title",
      "duties":     ["short duty 1", "short duty 2"],
      "start_date": "MM/YYYY",
      "end_date":   "MM/YYYY or present"
    }}
  ],
  "education": {{
    "higher_education": [
      {{"years": "YYYY-YYYY", "institution": "Name", "field": "Field"}}
    ],
    "further_training": [
      {{"years": "YYYY-YYYY", "institution": "Provider", "field": "Course"}}
    ]
  }},
  "language_skills": [
    {{"language": "German", "level": "B2 / Fluent / Native"}}
  ],
  "profile_summary": "~150 German words professional summary"
}}

Rules:
- employment_history: newest first
- duties: array of SHORT strings, NOT one long string
- Missing fields: null for strings, [] for arrays
- profile_summary: write in German, ~150 words

CV TEXT:
{text}
"""


# ── Robust JSON parser ────────────────────────────────────────────────────────

def _parse_json(raw: str) -> dict:
    """
    Parse JSON from any LLM response robustly:
    - strips ```json ... ``` and ``` ... ``` fences
    - extracts the first {...} block if surrounded by other text
    """
    if not raw:
        raise ValueError("Empty response from AI provider.")

    # 1. Strip common markdown fences
    cleaned = raw.strip()
    # Remove ```json...``` or ```...```
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\s*```\s*$", "", cleaned, flags=re.MULTILINE)
    cleaned = cleaned.strip()

    # 2. Try direct parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # 3. Extract first { ... } block
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except json.JSONDecodeError:
            pass

    raise ValueError(
        f"Could not parse JSON from AI response. Raw response preview:\n{raw[:500]}"
    )


# ── OpenAI-compatible helper (reused by many providers) ──────────────────────

def _openai_compat(text: str, model: str, api_key: str, base_url: str) -> dict:
    """Generic call using the standard CV build_prompt()."""
    return _openai_compat_direct(build_prompt(text), model, api_key, base_url)


def _openai_compat_direct(prompt: str, model: str, api_key: str, base_url: str) -> dict:
    """Generic call that accepts a pre-built prompt string."""
    from openai import OpenAI
    client = OpenAI(api_key=api_key, base_url=base_url)
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_MSG},
                {"role": "user",   "content": prompt},
            ],
            response_format={"type": "json_object"},
        )
    except Exception as e:
        # Fallback to older response format (non-JSON) if the model doesn't support it or if it's not the latest SDK
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_MSG},
                    {"role": "user",   "content": prompt},
                ],
            )
        except Exception as inner_e:
            masked = f"{api_key[:4]}...{api_key[-4:]}" if api_key and len(api_key) > 8 else "****"
            raise RuntimeError(f"AI Provider ({base_url}) error with key {masked}: {inner_e}") from inner_e
    return _parse_json(resp.choices[0].message.content)


# ── Provider-specific callers ─────────────────────────────────────────────────

def _call_openai(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://api.openai.com/v1")


def _call_gemini(text: str, model: str, api_key: str) -> dict:
    from google import genai
    from google.genai import types
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(
        model=model,
        contents=SYSTEM_MSG + "\n\n" + build_prompt(text),
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )
    return _parse_json(resp.text)


def _call_anthropic(text: str, model: str, api_key: str) -> dict:
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=model,
        max_tokens=4096,   # IMPORTANT: was 1024, too small for a full CV
        system=SYSTEM_MSG,
        messages=[{"role": "user", "content": build_prompt(text)}],
    )
    return _parse_json(msg.content[0].text)


def _call_mistral(text: str, model: str, api_key: str) -> dict:
    from mistralai import Mistral
    client = Mistral(api_key=api_key)
    resp = client.chat.complete(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_MSG},
            {"role": "user",   "content": build_prompt(text)},
        ],
        response_format={"type": "json_object"},
    )
    return _parse_json(resp.choices[0].message.content)


def _call_deepseek(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://api.deepseek.com/v1")


def _call_grok(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://api.x.ai/v1")


def _call_kimi(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://api.moonshot.cn/v1")


def _call_qwen(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://dashscope.aliyuncs.com/compatible-mode/v1")


def _call_perplexity(text: str, model: str, api_key: str) -> dict:
    return _openai_compat(text, model, api_key, "https://api.perplexity.ai/v1")


def _call_ollama(text: str, model: str, api_key: str) -> dict:
    host = api_key.rstrip("/") if api_key.startswith("http") else "http://localhost:11434"
    return _openai_compat(text, model, "ollama", f"{host}/v1")


# ── Dispatch table ────────────────────────────────────────────────────────────

DISPATCH = {
    "OpenAI":     _call_openai,
    "Gemini":     _call_gemini,
    "Anthropic":  _call_anthropic,
    "Mistral":    _call_mistral,
    "DeepSeek":   _call_deepseek,
    "Grok (xAI)": _call_grok,
    "Kimi K2":    _call_kimi,
    "Qwen":       _call_qwen,
    "Perplexity": _call_perplexity,
    "Ollama":     _call_ollama,
}


# ── Public API ────────────────────────────────────────────────────────────────

def get_cv_data(
    file_path: str,
    provider: str = "OpenAI",
    model: str = "gpt-4o-mini",
    api_key: str | None = None,
) -> dict:
    """Extract structured CV data using the requested AI provider."""
    # Use extract_text_from_any to support scanned PDF fallback
    text = extract_text_from_any(file_path, provider=provider, model=model, api_key=api_key or "")
    return DISPATCH[provider](text, model, api_key or "")


# ── Identcheck prompt ─────────────────────────────────────────────────────────

def build_identity_prompt(text: str) -> str:
    return f"""You are an identity document parser for HR compliance checks.
Extract ALL identity-related fields from the document below.
Return ONLY a raw JSON object — no markdown fences, no explanation.

Required JSON structure:
{{
  "first_name":                "First name(s)",
  "last_name":                 "Last name / surname",
  "full_name":                 "Last Name, First Name",
  "birth_date":                "DD.MM.YYYY",
  "birth_place":               "City, Country",
  "nationality":               "Nationality",
  "gender":                    "Männlich | Weiblich | Divers | null",

  "document_type":             "Personalausweis | Reisepass | Aufenthaltstitel | Sonstige",
  "document_number":           "Document / Passport number",
  "document_issue_date":       "DD.MM.YYYY or null",
  "document_expiry_date":      "DD.MM.YYYY",
  "document_issuing_authority":"Issuing authority / country",

  "residence_permit_type":     "Permit type string or null",
  "residence_permit_number":   "Number or null",
  "residence_permit_expiry":   "DD.MM.YYYY or null",

  "work_permit":               "Ja | Nein | Unbeschränkt | Beschränkt | null",
  "work_permit_expiry":        "DD.MM.YYYY or null",

  "address":                   "Full address string or null",
  "phone":                     "Phone number or null",
  "email":                     "Email or null",

  "checked_by":                null,
  "check_date":                null,
  "notes":                     null
}}

Rules:
- Use null for any field not found in the document.
- Dates always in DD.MM.YYYY format.
- document_type must be one of the enum values listed.

Document text:
{text}
"""


def get_identity_data(
    file_path: str,
    provider: str = "OpenAI",
    model: str = "gpt-4.1-mini",
    api_key: str | None = None,
) -> dict:
    """Extract identity document fields for the Identcheck Vorlage."""
    # Use extract_text_from_any to support scanned PDF fallback
    text = extract_text_from_any(file_path, provider=provider, model=model, api_key=api_key or "")

    if provider not in DISPATCH:
        raise ValueError(f"Unknown provider '{provider}'. Choose from: {list(DISPATCH)}")

    if provider != "Ollama" and not api_key:
        raise ValueError(f"No API key provided for {provider}.")

    prompt = build_identity_prompt(text)
    api_key = api_key or ""

    # Call each provider directly with the identity prompt
    if provider == "OpenAI":
        return _openai_compat_direct(prompt, model, api_key, "https://api.openai.com/v1")
    elif provider == "Gemini":
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=api_key)
            resp = client.models.generate_content(
                model=model,
                contents=SYSTEM_MSG + "\n\n" + prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json"),
            )
            return _parse_json(resp.text)
        except Exception as _e:
            masked = f"{api_key[:4]}...{api_key[-4:]}" if api_key and len(api_key) > 8 else "****"
            raise RuntimeError(f"Gemini error with key {masked}: {_e}") from _e
    elif provider == "Anthropic":
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model=model, max_tokens=2048, system=SYSTEM_MSG,
                messages=[{"role": "user", "content": prompt}],
            )
            return _parse_json(msg.content[0].text)
        except Exception as _e:
            masked = f"{api_key[:4]}...{api_key[-4:]}" if api_key and len(api_key) > 8 else "****"
            raise RuntimeError(f"Anthropic error with key {masked}: {_e}") from _e
    elif provider == "Mistral":
        return _openai_compat_direct(prompt, model, api_key, "https://api.mistral.ai/v1")
    elif provider == "DeepSeek":
        return _openai_compat_direct(prompt, model, api_key, "https://api.deepseek.com/v1")
    elif provider == "Grok (xAI)":
        return _openai_compat_direct(prompt, model, api_key, "https://api.x.ai/v1")
    elif provider == "Kimi K2":
        return _openai_compat_direct(prompt, model, api_key, "https://api.moonshot.cn/v1")
    elif provider == "Qwen":
        return _openai_compat_direct(prompt, model, api_key, "https://dashscope.aliyuncs.com/compatible-mode/v1")
    elif provider == "Perplexity":
        return _openai_compat_direct(prompt, model, api_key, "https://api.perplexity.ai")
    elif provider == "Ollama":
        host = api_key.rstrip("/") if api_key.startswith("http") else "http://localhost:11434"
        return _openai_compat_direct(prompt, model, "ollama", f"{host}/v1")
    else:
        raise ValueError(f"Unknown provider: {provider}")


# ── Image → text extraction (for ID scans / JPG uploads) ─────────────────────

def extract_text_from_image_openai(image_path: str, api_key: str,
                                    base_url: str = "https://api.openai.com/v1",
                                    model: str = "gpt-4o") -> str:
    """Use OpenAI vision to OCR an image file."""
    import base64
    from openai import OpenAI

    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    ext = os.path.splitext(image_path)[1].lower().lstrip(".")
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"

    client = OpenAI(api_key=api_key, base_url=base_url)
    resp = client.chat.completions.create(
        model=model,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": (
                    "Extract ALL text from this image exactly as it appears. "
                    "Include all visible fields, numbers, dates and names. "
                    "Return plain text only."
                )},
                {"type": "image_url", "image_url": {
                    "url": f"data:{mime};base64,{b64}"
                }},
            ],
        }],
        max_tokens=2048,
    )
    return resp.choices[0].message.content or ""


def extract_text_from_image_gemini(image_path: str, api_key: str,
                                    model: str = "gemini-2.0-flash") -> str:
    """Use Gemini vision to OCR an image file."""
    import base64
    from google import genai
    from google.genai import types

    with open(image_path, "rb") as f:
        b64_bytes = f.read()

    ext = os.path.splitext(image_path)[1].lower().lstrip(".")
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"

    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(
        model=model,
        contents=[
            types.Part.from_bytes(data=b64_bytes, mime_type=mime),
            "Extract ALL text from this image exactly as it appears. Return plain text only.",
        ],
    )
    return resp.text or ""


def extract_text_from_any(file_path: str, provider: str = "OpenAI",
                           model: str = "gpt-4o", api_key: str = "") -> str:
    """
    Extract text from PDF, DOCX, or image (JPG/PNG/JPEG).
    For PDFs: tries native text extraction first; if empty (scanned PDF),
    falls back to vision OCR using the selected provider.
    Images always use vision OCR.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        if text.strip():
            return text
        # ── Scanned PDF fallback: use vision OCR ─────────────────────────
        if provider == "Gemini" and api_key:
            return extract_text_from_scanned_pdf_gemini(file_path, api_key, model)
        vision_model = "gpt-4o" if provider == "OpenAI" else model
        base_url = {
            "OpenAI":     "https://api.openai.com/v1",
            "Grok (xAI)": "https://api.x.ai/v1",
            "DeepSeek":   "https://api.deepseek.com/v1",
        }.get(provider, "https://api.openai.com/v1")
        if api_key:
            return extract_text_from_scanned_pdf_openai(file_path, api_key, base_url, vision_model)
        # No key available — raise a clear, actionable error
        raise ValueError(
            "This PDF appears to be a scanned (image-only) document with no text layer. "
            "To process it, please go to ⚙️ Settings and add an API key for GPT-4o (OpenAI) "
            "or Gemini — both support image OCR and can read scanned documents."
        )

    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
        if text.strip():
            return text
        # ── Scanned DOCX fallback: extract images and OCR them ───────────
        try:
            doc = Document(file_path)
            parts = []
            import tempfile
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                        tmp.write(img_data)
                        tmp_img_path = tmp.name
                    try:
                        if provider == "Gemini":
                            parts.append(extract_text_from_image_gemini(tmp_img_path, api_key, model))
                        else:
                            vision_model = "gpt-4o" if provider == "OpenAI" else model
                            base_url = {
                                "OpenAI":     "https://api.openai.com/v1",
                                "Grok (xAI)": "https://api.x.ai/v1",
                                "DeepSeek":   "https://api.deepseek.com/v1",
                            }.get(provider, "https://api.openai.com/v1")
                            parts.append(extract_text_from_image_openai(tmp_img_path, api_key, base_url, vision_model))
                    finally:
                        os.unlink(tmp_img_path)
            if parts:
                return "\n\n".join(parts)
        except Exception:
            pass
        return text # Still empty

    elif ext in (".jpg", ".jpeg", ".png", ".webp"):
        if provider == "Gemini":
            return extract_text_from_image_gemini(file_path, api_key, model)
        # Default: OpenAI vision (also works for DeepSeek/Grok that mirror OAI API)
        vision_model = "gpt-4o" if provider == "OpenAI" else model
        base_url = {
            "OpenAI":     "https://api.openai.com/v1",
            "Grok (xAI)": "https://api.x.ai/v1",
            "DeepSeek":   "https://api.deepseek.com/v1",
        }.get(provider, "https://api.openai.com/v1")
        return extract_text_from_image_openai(file_path, api_key, base_url, vision_model)
    else:
        raise ValueError(f"Unsupported file format for Lebenslauf extraction: {ext}")


# ── Lebenslauf: multi-file extraction ────────────────────────────────────────

def get_lebenslauf_data(
    file_paths: list[str],
    job_role: str = "Schweißer",
    provider: str = "OpenAI",
    model: str = "gpt-4o",
    api_key: str | None = None,
) -> dict:
    """
    Extract and restructure data from one or more documents (CV + ID + extras)
    into the Vorlage Lebenslauf JSON structure.

    Parameters
    ----------
    file_paths : list of absolute paths — CV first, then ID, then extras
    job_role   : trade for the role-aware prompt
    provider   : AI provider name (must be in DISPATCH)
    model      : model name
    api_key    : API key / Ollama host

    Returns
    -------
    dict — structured JSON matching the lebenslauf_prompts schema
    """
    from lebenslauf_prompts import build_lebenslauf_prompt, LEBENSLAUF_SYSTEM_MSG

    if not api_key and provider != "Ollama":
        raise ValueError(f"No API key provided for {provider}.")

    resolved_key = api_key or ""

    # ── Step 1: extract text from every file ─────────────────────────────────
    all_text_parts: list[str] = []
    for i, fp in enumerate(file_paths):
        label = "LEBENSLAUF" if i == 0 else f"DOKUMENT {i}"
        try:
            txt = extract_text_from_any(fp, provider=provider, model=model,
                                        api_key=resolved_key)
            if txt.strip():
                all_text_parts.append(f"=== {label} ===\n{txt.strip()}")
        except Exception as e:
            all_text_parts.append(f"=== {label} ===\n[Konnte nicht gelesen werden: {e}]")

    combined_text = "\n\n".join(all_text_parts)
    if not combined_text.strip():
        raise ValueError("Kein Text aus den hochgeladenen Dokumenten extrahierbar.")

    # ── Step 2: build prompt and call AI ─────────────────────────────────────
    prompt = build_lebenslauf_prompt(combined_text, job_role)

    if provider == "OpenAI":
        return _openai_compat_direct(prompt, model, resolved_key, "https://api.openai.com/v1")
    elif provider == "Gemini":
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=resolved_key)
        resp = client.models.generate_content(
            model=model,
            contents=LEBENSLAUF_SYSTEM_MSG + "\n\n" + prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json"),
        )
        return _parse_json(resp.text)
    elif provider == "Anthropic":
        import anthropic
        client = anthropic.Anthropic(api_key=resolved_key)
        msg = client.messages.create(
            model=model, max_tokens=4096,
            system=LEBENSLAUF_SYSTEM_MSG,
            messages=[{"role": "user", "content": prompt}],
        )
        return _parse_json(msg.content[0].text)
    elif provider == "Mistral":
        from mistralai import Mistral
        client = Mistral(api_key=resolved_key)
        resp = client.chat.complete(
            model=model,
            messages=[{"role": "system", "content": LEBENSLAUF_SYSTEM_MSG},
                      {"role": "user",   "content": prompt}],
            response_format={"type": "json_object"},
        )
        return _parse_json(resp.choices[0].message.content)
    elif provider == "DeepSeek":
        return _openai_compat_direct(prompt, model, resolved_key, "https://api.deepseek.com/v1")
    elif provider == "Grok (xAI)":
        return _openai_compat_direct(prompt, model, resolved_key, "https://api.x.ai/v1")
    elif provider == "Kimi K2":
        return _openai_compat_direct(prompt, model, resolved_key, "https://api.moonshot.cn/v1")
    elif provider == "Qwen":
        return _openai_compat_direct(prompt, model, resolved_key,
                                     "https://dashscope.aliyuncs.com/compatible-mode/v1")
    elif provider == "Perplexity":
        return _openai_compat_direct(prompt, model, resolved_key, "https://api.perplexity.ai")
    elif provider == "Ollama":
        host = resolved_key.rstrip("/") if resolved_key.startswith("http") else "http://localhost:11434"
        return _openai_compat_direct(prompt, model, "ollama", f"{host}/v1")
    else:
        raise ValueError(f"Unknown provider for Lebenslauf extraction: {provider}")

    # This line should never be reached due to the and 'else' raise above.
    return {}


# ── Key validation ─────────────────────────────────────────────────────────────

def validate_key(provider: str, api_key: str, model: str = "") -> tuple[bool, str]:
    """
    Lightweight ping to verify an API key is valid.
    Returns (success: bool, message: str).
    """
    try:
        if provider == "Gemini":
            from google import genai
            client = genai.Client(api_key=api_key)
            # Ask for a tiny response
            resp = client.models.generate_content(
                model=model or "gemini-2.0-flash",
                contents="Say OK",
            )
            _ = resp.text
            return True, "✅ Key is valid"

        elif provider == "Anthropic":
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model=model or "claude-3-haiku-20240307",
                max_tokens=5,
                messages=[{"role": "user", "content": "Say OK"}],
            )
            _ = msg.content[0].text
            return True, "✅ Key is valid"

        elif provider == "Mistral":
            from openai import OpenAI
            client = OpenAI(api_key=api_key, base_url="https://api.mistral.ai/v1")
            client.models.list()
            return True, "✅ Key is valid"

        elif provider == "Ollama":
            import urllib.request
            host = api_key.rstrip("/") if api_key.startswith("http") else "http://localhost:11434"
            with urllib.request.urlopen(f"{host}/api/tags", timeout=5) as r:
                _ = r.read()
            return True, "✅ Ollama reachable"

        else:
            # OpenAI-compatible: OpenAI, DeepSeek, Grok, Kimi, Qwen, Perplexity
            BASE_URLS = {
                "OpenAI":     "https://api.openai.com/v1",
                "DeepSeek":   "https://api.deepseek.com/v1",
                "Grok (xAI)": "https://api.x.ai/v1",
                "Kimi K2":    "https://api.moonshot.cn/v1",
                "Qwen":       "https://dashscope.aliyuncs.com/compatible-mode/v1",
                "Perplexity": "https://api.perplexity.ai",
            }
            from openai import OpenAI
            base_url = BASE_URLS.get(provider, "https://api.openai.com/v1")
            client = OpenAI(api_key=api_key, base_url=base_url)
            models_list = client.models.list()
            return True, f"✅ Key is valid ({len(list(models_list.data))} models)"

    except Exception as e:
        msg = str(e)
        if "401" in msg or "authentication" in msg.lower() or "invalid" in msg.lower():
            return False, f"❌ Invalid key: {msg[:120]}"
        return False, f"❌ Error: {msg[:120]}"


if __name__ == "__main__":
    pass




