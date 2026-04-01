from docx import Document
from docx.shared import Pt, Cm, RGBColor
import os

doc = Document()

# Add styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# 1. Kopfbereich
doc.add_heading('KANDIDATENPROFIL', 0)
doc.add_paragraph('Titel des Job Postings: {{ job_role }}')
doc.add_paragraph('Stundenverrechnungssatz (SVS): €')
doc.add_paragraph('Möglicher Starttermin (verfügbar ab / Kündigungsfrist): {{ starttermin }}')
doc.add_paragraph('')
doc.add_heading('Zusätzliche Informationen zum Profil:', level=2)
doc.add_paragraph('{{ zusammenfassung }}')
doc.add_paragraph('')

# 2. Berufserfahrung
doc.add_heading('Berufserfahrung (Angaben bis zum aktuellen Zeitpunkt, mind. Monatsangaben):', level=1)
# Create a table for Berufserfahrung exactly as requested: Von- bis | Firma ...
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Von - bis'
hdr_cells[1].text = 'Firma'
row2 = table.rows[1].cells
row2[0].text = 'Bezeichnung / Tätigkeit'
row2[1].text = 'Prägnante Aufgabenbeschreibung'

# Add tag rows using run placeholders so it deletes correctly if empty
doc.add_paragraph('')
doc.add_paragraph('{%p for job in berufserfahrung %}', style='Normal')
t = doc.add_table(rows=2, cols=2)
t.style = 'Table Grid'
c1 = t.rows[0].cells
c1[0].text = '{{ job.von }} - {{ job.bis }}'
c1[1].text = '{{ job.arbeitgeber }}'
c2 = t.rows[1].cells
c2[0].text = '{{ job.position }}'
c2[1].text = '{{ job.taetigkeiten }}'
doc.add_paragraph('{%p endfor %}', style='Normal')
doc.add_paragraph('')

# 3. Ausbildung
doc.add_heading('Ausbildung:', level=1)
doc.add_paragraph('{%p for b in bildung %}')
tb = doc.add_table(rows=3, cols=2)
tb.style = 'Table Grid'
tb.rows[0].cells[0].text = 'Von- bis'
tb.rows[0].cells[1].text = '{{ b.jahre }}'

tb.rows[1].cells[0].text = 'Name der Hochschule/ Ausbildungsbetrieb'
tb.rows[1].cells[1].text = '{{ b.einrichtung }}'

tb.rows[2].cells[0].text = 'Abschluss'
tb.rows[2].cells[1].text = '{{ b.abschluss }}'
doc.add_paragraph('{%p endfor %}')
doc.add_paragraph('')

# 4. Weiterbildung
doc.add_heading('Weiterbildung:', level=1)
doc.add_paragraph('{%p for w in weiterbildung %}')
tw = doc.add_table(rows=2, cols=2)
tw.style = 'Table Grid'
tw.rows[0].cells[0].text = 'Von- bis'
tw.rows[0].cells[1].text = '{{ w.jahre }}'

tw.rows[1].cells[0].text = 'Name der Weiterbildungsstätte / Kurs'
tw.rows[1].cells[1].text = '{{ w.anbieter }} - {{ w.kurs }}'
doc.add_paragraph('{%p endfor %}')
doc.add_paragraph('')

# 5. Zertifikate
doc.add_heading('vorhandene Zertifikate/ Nachweise:', level=1)
doc.add_paragraph('{%p for z in zertifikate %}')
doc.add_paragraph('- {{ z.bezeichnung }} (gültig bis: {{ z.gueltig_bis }})')
doc.add_paragraph('{%p endfor %}')
doc.add_paragraph('')

# 6. Kompetenzen & Kenntnisse
doc.add_heading('Kompetenzen:', level=1)
doc.add_paragraph('Erfahrungen:')
doc.add_paragraph('Projekterfahrung:\tKeine')
doc.add_paragraph('Führungserfahrung:\tKeine')
doc.add_paragraph('')
doc.add_heading('EDV-Kenntnisse:', level=2)
doc.add_paragraph('MS-Word:\t\tKeine')
doc.add_paragraph('MS-EXCEL:\t\tKeine')
doc.add_paragraph('MS-PowerPoint:\tKeine')
doc.add_paragraph('MS-Outlook:\t\tKeine')
doc.add_paragraph('MS-Access:\t\tKeine')
doc.add_paragraph('SAP:\t\tKeine')
doc.add_paragraph('Weitere:')
doc.add_paragraph('')

# 7. Sonstige Techniken (Fähigkeiten)
doc.add_heading('Sonstige Techniken:', level=1)
doc.add_paragraph('{%p for f in faehigkeiten %}')
tf = doc.add_table(rows=1, cols=2)
tf.rows[0].cells[0].text = '{{ f.name }}:'
tf.rows[0].cells[1].text = '{{ f.beschreibung }}'
doc.add_paragraph('{%p endfor %}')
doc.add_paragraph('')

# 8. Sprachkenntnisse
doc.add_heading('Sprachkenntnisse:', level=1)
doc.add_paragraph('{%p for s in sprachen %}')
ts = doc.add_table(rows=1, cols=2)
ts.rows[0].cells[0].text = '{{ s.sprache }}:'
ts.rows[0].cells[1].text = '{{ s.niveau }}'
doc.add_paragraph('{%p endfor %}')
doc.add_paragraph('')

# Save to templates directory
out_dir = os.path.join(os.path.dirname(__file__), "templates")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "Vorlage_Lebenslauf_Kandidatenprofil_V2_Mit_Tags.docx")
doc.save(out_path)
print(f"Created template at {out_path}")
