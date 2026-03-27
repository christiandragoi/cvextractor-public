"""
lebenslauf_builder.py
=====================
Builds a styled German Word document (Vorlage Lebenslauf) from structured
JSON data produced by the AI extractor.

Layout mirrors the user-specified template:
  - Header with company logo placeholder + KANDIDATENPROFIL title
  - Kandidatendaten section
  - Berufserfahrung  (flowing paragraphs — date line, bold labels, details)
  - Bildungseinrichtung
  - Weiterbildung
  - Zertifikate
  - Fähigkeiten    (key: value pairs, bold keys)
  - Sprachkenntnisse (simple label: value)
  - Profil-Zusammenfassung
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Font used everywhere
FONT_NAME = "Calibri"
FONT_SIZE = 11  # pt — default body size

# ---------------------------------------------------------------------------
# Colour palette (matching a professional blue/dark theme)
# ---------------------------------------------------------------------------
DARK_BLUE   = RGBColor(0x0F, 0x34, 0x60)
MID_BLUE    = RGBColor(0x1E, 0x40, 0xAF)
LIGHT_GREY  = RGBColor(0xF1, 0xF5, 0xF9)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x0A, 0x0A, 0x0A)
TEXT_GREY   = RGBColor(0x47, 0x55, 0x69)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, rgb_hex: str):
    """Set table cell background colour using raw XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), rgb_hex)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def _para_border_bottom(para):
    """Add a thin bottom border to a paragraph (section divider)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1E40AF")
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_run_font(run, size=FONT_SIZE, bold=False, italic=False, color=BLACK):
    """Apply consistent Calibri formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1):
    """Add a styled section heading with bottom border."""
    p = doc.add_paragraph()
    _para_border_bottom(p)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    _set_run_font(run, size=FONT_SIZE, bold=True, color=MID_BLUE)
    return p


def _normal(doc: Document, text: str, bold: bool = False, italic: bool = False,
            size: int = FONT_SIZE, color: RGBColor = BLACK, indent: float = 0.0):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _safe(data: dict, key: str, default: str = "") -> str:
    val = data.get(key, default)
    if val is None:
        return default
    return str(val).strip() or default


def _safe_list(data: dict, key: str):
    val = data.get(key, [])
    return val if isinstance(val, list) else []


# ---------------------------------------------------------------------------
# Public builder
# ---------------------------------------------------------------------------

def build_lebenslauf_docx(data: dict, job_role: str = "Schweißer",
                           output_path: str | None = None) -> str:
    """
    Build a Vorlage-style Lebenslauf Word document from structured JSON data.

    Parameters
    ----------
    data        : dict — structured JSON from AI extraction
    job_role    : str  — used in the document title
    output_path : str  — where to save the .docx; if None, uses /tmp

    Returns
    -------
    str — absolute path to the generated .docx file
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Default paragraph style — Calibri 11 ───────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)

    # ════════════════════════════════════════════════════════════════════════
    # HEADER — Title + Kopfbereich Table
    # ════════════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("K A N D I D A T E N P R O F I L")
    _set_run_font(title_run, size=18, bold=True, color=DARK_BLUE)

    doc.add_paragraph()

    # Kopfbereich Table (as requested by user)
    kb_tbl = doc.add_table(rows=4, cols=3)
    kb_tbl.style = "Table Grid"
    kb_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header labels
    kb_labels = [
        "Titel des Job Postings",
        "Einkaufskurzprofil (EKP)",
        "Stundenverrechnungssatz (SVS)",
        "Möglicher Starttermin"
    ]
    
    # EKP and SVS placeolders
    ekp_val = "X|YYY|XXX|Z"
    svs_val = "€"
    start_val = data.get("starttermin", "01.03.2026") # Fallback to user example date if not found
    
    for i, label in enumerate(kb_labels):
        row = kb_tbl.rows[i]
        # Label cell
        _set_cell_bg(row.cells[0], "F1F5F9")
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        _set_run_font(lr, size=10, bold=True)
        
        # Value cells (merged or separate)
        curr_val = ""
        if i == 0: curr_val = "" # Job Title empty as requested
        elif i == 1: curr_val = ekp_val
        elif i == 2: curr_val = svs_val
        elif i == 3: curr_val = start_val
        
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(curr_val)
        _set_run_font(vr, size=10)
        
        vp2 = row.cells[2].paragraphs[0]
        vr2 = vp2.add_run(curr_val)
        _set_run_font(vr2, size=10)

    doc.add_paragraph() # spacer

    # Verleiher header (dark bar)
    verleiher_tbl = doc.add_table(rows=1, cols=2)
    verleiher_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    verleiher_tbl.style = "Table Grid"
    _set_cell_bg(verleiher_tbl.cell(0, 0), "0F3460")
    _set_cell_bg(verleiher_tbl.cell(0, 1), "0F3460")
    
    lc = verleiher_tbl.cell(0, 0); rc = verleiher_tbl.cell(0, 1)
    lc.width = Cm(9); rc.width = Cm(9)
    lc_p = lc.paragraphs[0]
    lr = lc_p.add_run("CD International GmbH\nPersonaldienstleistung")
    _set_run_font(lr, size=10, bold=True, color=WHITE)
    
    rc_p = rc.paragraphs[0]; rc_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rc_p.add_run(
        "Ansprechpartner: [Name Verleiher]\n"
        "Tel: +49 [xxx] [xxx]\n"
        "E-Mail: info@cd-international.de"
    )
    _set_run_font(rr, size=9, color=WHITE)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PERSÖNLICHE DATEN des Zeitarbeitnehmers
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Persönliche Daten des Zeitarbeitnehmers")

    fields = [
        ("Anrede",             _safe(data, "anrede")),
        ("Vorname",            _safe(data, "vorname")),
        ("Nachname",           _safe(data, "nachname")),
        ("Geburtsdatum",       _safe(data, "geburtsdatum")),
        ("Geburtsort",         _safe(data, "geburtsort")),
        ("Staatsangehörigkeit",_safe(data, "staatsangehoerigkeit")),
    ]
    kd_tbl = doc.add_table(rows=len(fields), cols=2)
    kd_tbl.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        row = kd_tbl.rows[i]
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(13)
        _set_cell_bg(row.cells[0], "EFF6FF")
        lp = row.cells[0].paragraphs[0]
        lr2 = lp.add_run(label)
        _set_run_font(lr2, size=FONT_SIZE, bold=True)
        rp = row.cells[1].paragraphs[0]
        vr = rp.add_run(value or "—")
        _set_run_font(vr, size=FONT_SIZE)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # BERUFSERFAHRUNG — flowing paragraphs (NO table)
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Berufserfahrung")

    experience = _safe_list(data, "berufserfahrung")
    if not experience:
        _normal(doc, "Keine Berufserfahrungsdaten gefunden.", italic=True, color=TEXT_GREY)
    else:
        # User wants chronological table style: Date Range Left, Details Right
        exp_tbl = doc.add_table(rows=0, cols=2)
        exp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for job in experience:
            row = exp_tbl.add_row()
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(14)
            
            # ── Date Range (Left) ──
            von = _safe(job, "von")
            bis = _safe(job, "bis")
            date_str = f"{von} – {bis}" if von else bis
            dp = row.cells[0].paragraphs[0]
            dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            dr = dp.add_run(date_str)
            _set_run_font(dr, size=FONT_SIZE, bold=True)
            
            # ── Job Details (Right) ──
            rp = row.cells[1].paragraphs[0]
            rp.paragraph_format.space_after = Pt(6)
            
            # Arbeitgeber / Land
            ag = _safe(job, "arbeitgeber")
            ort_land = _safe(job, "ort_land")
            ag_line = ag
            if ort_land and ort_land.upper() not in ag.upper():
                ag_line += f", {ort_land}"
            
            pos = _safe(job, "position")
            tat = _safe(job, "taetigkeiten")
            
            r_ag = rp.add_run(f"Arbeitgeber: {ag_line}\n")
            _set_run_font(r_ag, size=FONT_SIZE, bold=True)
            
            r_pos = rp.add_run(f"Position: {pos}\n")
            _set_run_font(r_pos, size=FONT_SIZE, bold=True)
            
            r_tat_lbl = rp.add_run("Tätigkeit: ")
            _set_run_font(r_tat_lbl, size=FONT_SIZE, bold=True)
            r_tat_val = rp.add_run(tat)
            _set_run_font(r_tat_val, size=FONT_SIZE)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # BILDUNGSEINRICHTUNG
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Bildungseinrichtung")

    bildung = _safe_list(data, "bildung")
    if not bildung:
        _normal(doc, "Keine Bildungsdaten gefunden.", italic=True, color=TEXT_GREY)
    else:
        for b in bildung:
            # Date line
            p_bj = doc.add_paragraph()
            p_bj.paragraph_format.space_before = Pt(6)
            p_bj.paragraph_format.space_after  = Pt(2)
            jr = p_bj.add_run(f"{_safe(b, 'jahre')}:")
            _set_run_font(jr, size=FONT_SIZE, bold=True, color=BLACK)

            # Institution + details (indented)
            einr = _safe(b, "einrichtung")
            abschluss = _safe(b, "abschluss")
            detail = einr
            if abschluss:
                detail += f" – {abschluss}"
            p_bd = doc.add_paragraph()
            p_bd.paragraph_format.space_before = Pt(0)
            p_bd.paragraph_format.space_after  = Pt(2)
            p_bd.paragraph_format.left_indent = Cm(2.0)
            bdr = p_bd.add_run(detail)
            _set_run_font(bdr, size=FONT_SIZE, color=BLACK)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # WEITERBILDUNG
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Weiterbildung")

    weiter = _safe_list(data, "weiterbildung")
    if not weiter:
        _normal(doc, "Keine Weiterbildungsdaten gefunden.", italic=True, color=TEXT_GREY)
    else:
        for w in weiter:
            # Date line
            p_wj = doc.add_paragraph()
            p_wj.paragraph_format.space_before = Pt(6)
            p_wj.paragraph_format.space_after  = Pt(2)
            wr_j = p_wj.add_run(f"{_safe(w, 'jahre')}:")
            _set_run_font(wr_j, size=FONT_SIZE, bold=True, color=BLACK)

            # Course + provider (indented)
            kurs = _safe(w, "kurs")
            anbieter = _safe(w, "anbieter")
            detail = anbieter
            if kurs:
                detail += f" – {kurs}"
            p_wd = doc.add_paragraph()
            p_wd.paragraph_format.space_before = Pt(0)
            p_wd.paragraph_format.space_after  = Pt(2)
            p_wd.paragraph_format.left_indent = Cm(2.0)
            wdr = p_wd.add_run(detail)
            _set_run_font(wdr, size=FONT_SIZE, color=BLACK)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # ZERTIFIKATE
    # ════════════════════════════════════════════════════════════════════════
    certs = _safe_list(data, "zertifikate")
    if certs:
        _heading(doc, "Zertifikate & Berechtigungen")
        for c in certs:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.8)
            label = _safe(c, "bezeichnung")
            aus   = _safe(c, "ausgestellt")
            bis   = _safe(c, "gueltig_bis")
            detail = ""
            if aus:  detail += f"  ausgestellt: {aus}"
            if bis and bis.lower() not in ("null", ""):
                detail += f"  /  gültig bis: {bis}"
            cr = p.add_run(label)
            _set_run_font(cr, size=FONT_SIZE, bold=True)
            if detail:
                dr2 = p.add_run(detail)
                _set_run_font(dr2, size=FONT_SIZE, color=TEXT_GREY)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # FÄHIGKEITEN — structured key: value pairs
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Fähigkeiten")

    faehig = data.get("faehigkeiten")
    if faehig:
        # Handle both old freetext format and new structured list format
        if isinstance(faehig, list):
            # New format: list of {"name": "...", "beschreibung": "..."}
            for item in faehig:
                if isinstance(item, dict):
                    name = _safe(item, "name")
                    beschr = _safe(item, "beschreibung")
                    p_fh = doc.add_paragraph()
                    p_fh.paragraph_format.space_before = Pt(0)
                    p_fh.paragraph_format.space_after  = Pt(2)
                    r_name = p_fh.add_run(f"{name}: ")
                    _set_run_font(r_name, size=FONT_SIZE, bold=True, color=BLACK)
                    r_desc = p_fh.add_run(beschr)
                    _set_run_font(r_desc, size=FONT_SIZE, bold=False, color=BLACK)
                elif isinstance(item, str):
                    _normal(doc, item, size=FONT_SIZE)
        elif isinstance(faehig, str):
            # Legacy freetext — just dump it
            for line in faehig.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Try to split on colon for key:value
                if ":" in line:
                    key, _, val = line.partition(":")
                    p_fh = doc.add_paragraph()
                    p_fh.paragraph_format.space_before = Pt(0)
                    p_fh.paragraph_format.space_after  = Pt(2)
                    r_key = p_fh.add_run(f"{key.strip()}: ")
                    _set_run_font(r_key, size=FONT_SIZE, bold=True, color=BLACK)
                    r_val = p_fh.add_run(val.strip())
                    _set_run_font(r_val, size=FONT_SIZE, bold=False, color=BLACK)
                else:
                    _normal(doc, line, size=FONT_SIZE)
    else:
        _normal(doc, "Keine Fähigkeiten angegeben.", italic=True, color=TEXT_GREY)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SPRACHKENNTNISSE — simple label: value (no table)
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Sprachkenntnisse")

    sprachen = _safe_list(data, "sprachen")
    if not sprachen:
        _normal(doc, "Keine Sprachkenntnisse angegeben.", italic=True, color=TEXT_GREY)
    else:
        for s in sprachen:
            sprache = _safe(s, "sprache")
            niveau  = _safe(s, "niveau")
            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_before = Pt(0)
            p_sp.paragraph_format.space_after  = Pt(2)
            r_sp = p_sp.add_run(f"{sprache}:\t")
            _set_run_font(r_sp, size=FONT_SIZE, bold=True, color=BLACK)
            r_nv = p_sp.add_run(niveau)
            _set_run_font(r_nv, size=FONT_SIZE, bold=False, color=BLACK)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # ZUSAMMENFASSUNG
    # ════════════════════════════════════════════════════════════════════════
    zusammen = _safe(data, "zusammenfassung")
    if zusammen:
        _heading(doc, "Profil-Zusammenfassung")
        p_zf = doc.add_paragraph()
        p_zf.paragraph_format.space_before = Pt(4)
        p_zf.paragraph_format.space_after  = Pt(4)
        zr = p_zf.add_run(zusammen)
        _set_run_font(zr, size=FONT_SIZE, color=TEXT_GREY)

    # ════════════════════════════════════════════════════════════════════════
    # Save
    # ════════════════════════════════════════════════════════════════════════
    if output_path is None:
        import tempfile
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"Lebenslauf_{_safe(data, 'nachname', 'Kandidat')}_{job_role}.docx"
        )
    doc.save(output_path)
    return output_path
