import io
import os
import re
import copy
from pathlib import Path
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.table import Table
except Exception:
    Document = None

logger = __import__("logging").getLogger(__name__)


class TemplatePopulationService:
    """
    Populate a DOCX template with structured candidate data.
    Supports:
      - Simple text placeholders: {{ first_name }}, {{ last_name }}, etc.
      - Repeated table rows for employment history with {{ job.xxx }} and {{ d }}
    """

    def __init__(self, template_path: str):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        if Document is None:
            raise RuntimeError("python-docx is not installed")
        self.doc = Document(str(self.template_path))

    def save(self, data: Dict[str, Any], output_path: str) -> str:
        """Populate template and save to output_path. Returns output_path."""
        self._populate(data)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        logger.info(f"Populated template saved to: {output_path}")
        return str(output_path)

    def _populate(self, data: Dict[str, Any]):
        """Main population logic."""
        self._replace_simple_fields(data)
        self._replace_job_rows(data)
        self._cleanup_remaining_placeholders()

    def _resolve_value(self, key: str, data: Dict[str, Any]) -> str:
        """Resolve a placeholder key to a string value."""
        if key in data and data[key] is not None:
            val = data[key]
            if isinstance(val, list):
                return ", ".join(str(v) for v in val)
            return str(val)

        if "." in key:
            parts = key.split(".")
            val = data
            for part in parts:
                if isinstance(val, dict) and part in val:
                    val = val[part]
                else:
                    val = None
                    break
            if val is not None:
                if isinstance(val, list):
                    return ", ".join(str(v) for v in val)
                return str(val)

        if key == "employment_history" and "jobs" in data:
            jobs = data["jobs"]
            if isinstance(jobs, list):
                return self._format_jobs(jobs)

        aliases = {
            "full_name": ["name", "candidate_name", "person_name"],
            "first_name": ["fname", "vorname"],
            "last_name": ["surname", "family_name", "nachname"],
            "date_of_birth": ["dob", "birth_date", "geburtstag", "geboren_am"],
            "place_of_birth": ["birth_place", "gebortsort"],
            "nationality": ["citizenship", "staatsangehoerigkeit", "nationalitaet"],
            "email": ["e_mail", "mail"],
            "phone": ["telefon", "tel", "mobile", "handynummer"],
            "address": ["street", "strasse", "wohnort"],
            "skills": ["faehigkeiten", "kentnisse"],
            "languages": ["sprachen"],
        }
        if key in aliases:
            for alt in aliases[key]:
                if alt in data and data[alt] is not None:
                    return str(data[alt])

        if "candidate" in data and isinstance(data["candidate"], dict):
            cand = data["candidate"]
            if key in cand and cand[key] is not None:
                return str(cand[key])
            for alt in aliases.get(key, []):
                if alt in cand and cand[alt] is not None:
                    return str(cand[alt])

        # Derive first_name / last_name from full_name
        if key in ("first_name", "last_name"):
            full_name = data.get("full_name") or data.get("candidate", {}).get("full_name", "")
            if full_name and isinstance(full_name, str):
                parts = full_name.strip().split()
                if key == "first_name":
                    return parts[0] if parts else ""
                else:
                    return " ".join(parts[1:]) if len(parts) > 1 else ""

        return ""

    def _format_jobs(self, jobs: List[Dict]) -> str:
        lines = []
        for j in jobs:
            line = f"{j.get('job_title','')} at {j.get('company','')} ({j.get('start_date','')} - {j.get('end_date','')})"
            lines.append(line)
        return "\n".join(lines)

    def _replace_simple_fields(self, data: Dict[str, Any]):
        """Replace {{ key }} placeholders in paragraphs and table cells."""
        placeholder_re = re.compile(r"\{\{\s*([a-zA-Z0-9_\.]+)\s*\}\}")
        for paragraph in self._iter_all_paragraphs():
            self._replace_in_paragraph(paragraph, placeholder_re, data)

    def _iter_all_paragraphs(self):
        for paragraph in self.doc.paragraphs:
            yield paragraph
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    def _replace_in_paragraph(self, paragraph, placeholder_re, data):
        """Replace placeholders across all runs in a paragraph by merging text first."""
        runs = paragraph.runs
        if not runs:
            return

        # Merge all run texts
        merged = "".join(r.text for r in runs)
        if not placeholder_re.search(merged):
            return

        def replacer(match):
            key = match.group(1).strip()
            if key.startswith("job.") or key == "d":
                return match.group(0)
            return self._resolve_value(key, data)

        new_merged = placeholder_re.sub(replacer, merged)
        if new_merged == merged:
            return

        # Write back: first run gets full text, rest cleared
        for i, run in enumerate(runs):
            run.text = new_merged if i == 0 else ""

    def _replace_job_rows(self, data: Dict[str, Any]):
        jobs = data.get("employment_history") or data.get("jobs") or data.get("candidate", {}).get("employment_history", [])
        if not isinstance(jobs, list) or not jobs:
            self._remove_empty_job_rows()
            return

        for table in self.doc.tables:
            self._process_job_table(table, jobs)

    def _process_job_table(self, table: Table, jobs: List[Dict]):
        job_fields = {"job.start_date", "job.end_date", "job.employer", "job.position",
                      "job.company", "job.title", "job.location", "job.description", "d"}

        template_row_idx = None
        template_row = None
        for idx, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            if any(f"{{{{ {f} }}}}" in row_text or f"{{{{{f}}}}}" in row_text for f in job_fields):
                template_row_idx = idx
                template_row = row
                break

        if template_row is None:
            return

        # Find the XML index of the template row (not the logical row index)
        xml_index = list(table._tbl).index(template_row._tr)

        # Build cloned rows for each job
        new_rows = []
        for job_index, job in enumerate(jobs):
            new_row = self._copy_table_row(template_row)
            if new_row is None:
                continue
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_job_placeholders_in_paragraph(paragraph, job, job_index + 1)
            new_rows.append(new_row)

        # Remove original template row
        template_row._tr.getparent().remove(template_row._tr)

        # Insert new rows at the XML position where the template row was
        for new_row in reversed(new_rows):
            table._tbl.insert(xml_index, new_row._tr)

    def _copy_table_row(self, row):
        try:
            new_tr = copy.deepcopy(row._tr)
            from docx.table import _Row
            return _Row(new_tr, row.table)
        except Exception as e:
            logger.warning(f"Could not copy table row: {e}")
            return None

    def _replace_job_placeholders_in_paragraph(self, paragraph, job: Dict, index: int):
        runs = paragraph.runs
        if not runs:
            return
        merged = "".join(r.text for r in runs)

        def replacer(match):
            key = match.group(1).strip()
            if key == "d":
                return str(index)
            if key.startswith("job."):
                field = key[4:]
                field_map = {
                    "start_date": ["start_date", "start", "von", "from"],
                    "end_date": ["end_date", "end", "bis", "to", "present", "heute"],
                    "employer": ["employer", "company", "firma", "arbeitgeber"],
                    "position": ["position", "job_title", "title", "rolle", "beruf", "stelle"],
                    "location": ["location", "ort", "standort"],
                    "description": ["description", "beschreibung", "taetigkeiten"],
                }
                candidates = field_map.get(field, [field])
                for c in candidates:
                    if c in job and job[c] is not None:
                        return str(job[c])
                # Fallback: if job_title contains both title and company, try to split
                if field in ("position", "employer") and "job_title" in job:
                    jt = str(job["job_title"])
                    # Try to split by common separators
                    if " bei " in jt.lower() or " at " in jt.lower():
                        parts = re.split(r'\s+(?:bei|at)\s+', jt, flags=re.IGNORECASE)
                        if len(parts) >= 2:
                            if field in ("position", "job_title", "title"):
                                return parts[0].strip()
                            else:
                                return parts[1].strip()
                    if " bei " in jt:
                        parts = jt.split(" bei ")
                        if len(parts) >= 2:
                            if field in ("position", "job_title", "title"):
                                return parts[0].strip()
                            else:
                                return parts[1].strip()
                return ""
            return match.group(0)

        new_merged = re.sub(r"\{\{\s*([a-zA-Z0-9_\.]+)\s*\}\}", replacer, merged)
        if new_merged == merged:
            return
        for i, run in enumerate(runs):
            run.text = new_merged if i == 0 else ""

    def _remove_empty_job_rows(self):
        job_pattern = re.compile(r"\{\{\s*(job\.[a-zA-Z0-9_]+|d)\s*\}\}")
        for table in self.doc.tables:
            rows_to_remove = []
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                if job_pattern.search(row_text) and not re.search(r"[A-Za-z0-9]{3,}", re.sub(r"\{\{.*?\}\}", "", row_text)):
                    rows_to_remove.append(row._tr)
            for tr in rows_to_remove:
                tr.getparent().remove(tr)

    def _cleanup_remaining_placeholders(self):
        placeholder_re = re.compile(r"\{\{\s*[a-zA-Z0-9_\.]+\s*\}\}")
        for paragraph in self._iter_all_paragraphs():
            runs = paragraph.runs
            if not runs:
                continue
            merged = "".join(r.text for r in runs)
            if not placeholder_re.search(merged):
                continue
            cleaned = placeholder_re.sub("", merged)
            for i, run in enumerate(runs):
                run.text = cleaned if i == 0 else ""


def extract_text_from_file(file_path: str) -> str:
    full_path = Path(file_path)
    if not full_path.exists():
        return ""
    suffix = full_path.suffix.lower()
    text = ""
    if suffix == ".docx" and Document:
        try:
            doc = Document(str(full_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
        except Exception:
            text = ""
    elif suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(full_path)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                text = "\n".join(pages)
        except Exception:
            text = ""
    elif suffix == ".doc":
        try:
            text = full_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = ""
    return text


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx" and Document:
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception:
            return ""
    elif suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                return "\n".join(pages)
        except Exception:
            return ""
    elif suffix == ".doc":
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    return ""
